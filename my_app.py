from docx import Document
from docx.shared import Inches


document = Document()

# profile picture

document.add_picture('me.jpg', width=Inches(2.0))

# name phone number and email details

name = input('what is your name?')

phone_number = input('what is your phone number?')

email = input('what is your email?')

document.add_paragraph( name + ' | ' + phone_number + ' | ' + email)

# about me

document.add_heading('about me')
about_me = input('Tell about yourself')
document.add_paragraph(about_me)

# work experience

document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From date')
to_date = input('To date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company)
p.add_run(experience_details)

# more experience

while True:
    has_more_experience = input('Do you have more experience? Yes or No ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company')
        from_date = input('From date')
        to_date = input('To date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break


#skills

document.add_heading('Skills')

skill=input('Enter skill')

p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill=input('Enter skill')

        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated"


document.save('cv.docx')


   

